"""
src/EEM_Project/costs/data_handler.py

Data input for costs (Block_TC and Block_TC_Gas).

Maps the wide cost layout into CostItem / CostProfile using the categories
defined in the core model:
  exploration, wells_capex, facilities_capex, opex, abandonment, other

Supports the same two-mode philosophy as production (xlsm import + JSON roundtrip).

This feeds the capital allowances module later.
"""

from __future__ import annotations

import io
import json
import tempfile
from pathlib import Path
from typing import Any, Optional, Union

import pandas as pd
from openpyxl import load_workbook

from EEM_Project.core.models import CostItem, CostProfile

UploadedFileLike = Union["BinaryIO", "UploadedFile"]  # Streamlit support

DEFAULT_XLSM = Path(r"C:\Users\Emmanuel Onwuka\Desktop\Acquisitions\Econ_Model_Draft_2.xlsm")


CATEGORY_MAP = {
    "Exploration": "exploration",
    "CAPEX Wells": "wells_capex",
    "CAPEX Facilities": "facilities_capex",
    "OPEX": "opex",
    # "Unit OPEX" is $/bbl — we can store as "other" or derive later; for now focus on $mm
    "Unit OPEX": "other",
}


def load_cost_profiles(
    xlsm_path: Optional[Path] = None,
    json_dir: Optional[Path] = None,
    fluid: str = "oil",   # "oil" or "gas"
) -> list[CostProfile]:
    if xlsm_path is None:
        xlsm_path = DEFAULT_XLSM

    if xlsm_path and xlsm_path.exists():
        return _load_costs_from_xlsm(xlsm_path, fluid=fluid)

    if json_dir is None:
        json_dir = Path("data/examples") / "costs"

    if json_dir.exists():
        return _load_costs_from_json(json_dir, fluid=fluid)

    return []


def _load_costs_from_xlsm(xlsm_path: Path, fluid: str = "oil") -> list[CostProfile]:
    sheet = "Block_TC" if fluid == "oil" else "Block_TC_Gas"
    print(f"[costs_handler] Loading costs from {sheet} in {xlsm_path}")

    wb = load_workbook(str(xlsm_path), data_only=True)
    ws = wb[sheet]

    # Discover blocks from row 1 (similar wide pattern, but 5 cost cols + separator)
    # Row 2 has the category headers
    blocks: list[tuple[int, str]] = []
    for c in range(3, ws.max_column + 1, 6):  # rough step; adjust if needed
        name = ws.cell(row=1, column=c).value
        if name and isinstance(name, str) and ("OML" in name or "FIELD" not in name.upper()):
            blocks.append((c, name.strip()))

    # Better discovery: look at row 2 categories to find start of each block group
    # Simpler: use the same stepping as production but account for 5 cols per field + separator
    # From inspection: groups of 5 useful cols (Exploration to OPEX), with separators.
    # For robustness we scan for known category in row 2.
    block_starts: list[tuple[int, str]] = []
    for c in range(2, ws.max_column + 1):
        cat = ws.cell(row=2, column=c).value
        if cat == "Exploration":
            # block name is typically in row 1, same group (a few cols before or at the Exploration col)
            for look in [0, -1, -2, 2, 3]:
                candidate = ws.cell(row=1, column=max(1, c + look)).value
                if candidate and isinstance(candidate, str) and ("OML" in candidate or "FIELD" not in candidate.upper()):
                    block_starts.append((c, candidate.strip()))
                    break

    if not block_starts:
        # fallback
        for c in range(3, ws.max_column + 1, 6):
            nm = ws.cell(row=1, column=c).value
            if nm and isinstance(nm, str) and "OML" in nm:
                block_starts.append((c, nm.strip()))

    years = []
    for r in range(4, ws.max_row + 1):
        y = ws.cell(row=r, column=1).value
        if isinstance(y, (int, float)) and 2000 < y < 2100:
            years.append(int(y))

    profiles: list[CostProfile] = []

    for start_col, name in block_starts:
        items: list[CostItem] = []
        for r, year in enumerate(years, start=4):
            for offset, excel_cat in enumerate(["Exploration", "CAPEX Wells", "CAPEX Facilities", "OPEX"]):
                col = start_col + offset
                amt = ws.cell(row=r, column=col).value
                if amt is None:
                    amt = 0.0
                try:
                    amt = float(amt)
                except Exception:
                    amt = 0.0

                model_cat = CATEGORY_MAP.get(excel_cat, "other")
                item = CostItem(
                    block_name=name,
                    year=year,
                    category=model_cat,  # type: ignore[arg-type]
                    amount_usd_mln=amt,
                    is_oil=(fluid == "oil"),
                    is_gas=(fluid == "gas"),
                )
                items.append(item)

        prof = CostProfile(block_name=name, items=items)
        profiles.append(prof)

    wb.close()
    print(f"[costs_handler] Loaded {len(profiles)} cost profiles for {fluid} ({len(years)} years each).")
    return profiles


def _load_costs_from_json(json_dir: Path, fluid: str = "oil") -> list[CostProfile]:
    profiles = []
    suffix = f"_{fluid}_costs.json" if fluid else ".json"
    for jf in sorted(json_dir.glob(f"*{suffix}")):
        with open(jf, "r", encoding="utf-8") as f:
            data = json.load(f)
        prof = CostProfile(**data)
        profiles.append(prof)
    return profiles


def save_cost_profiles(profiles: list[CostProfile], out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    for prof in profiles:
        safe = prof.block_name.replace(" ", "_").replace("/", "_")
        with open(out_dir / f"{safe}_costs.json", "w", encoding="utf-8") as f:
            json.dump(prof.model_dump(), f, indent=2, default=str)
    print(f"[costs_handler] Saved {len(profiles)} cost profiles to {out_dir}")


# ============================================================
# Flexible multi-format import for costs (parallel to production)
# ============================================================

def _cost_items_from_dataframe(
    df: pd.DataFrame, column_map: Optional[dict[str, str]] = None, **kwargs
) -> list[CostItem]:
    """Convert generic DF to CostItems. Expects block_name, year, category, amount_usd_mln, and optional is_oil/is_gas or fluid."""
    if column_map:
        df = df.rename(columns=column_map)

    # Aliases
    aliases = {
        "block_name": ["block_name", "block", "field", "well", "name"],
        "year": ["year", "yr"],
        "category": ["category", "cost_type", "type"],
        "amount_usd_mln": ["amount_usd_mln", "amount", "value", "cost", "$mm"],
        "is_oil": ["is_oil", "oil", "fluid_oil"],
        "is_gas": ["is_gas", "gas", "fluid_gas"],
        "fluid_type": ["fluid_type", "fluid", "type"],
    }
    col_map = {}
    for std, poss in aliases.items():
        for col in df.columns:
            if str(col).lower().strip() in poss:
                col_map[col] = std
                break
    df = df.rename(columns=col_map)

    required = ["block_name", "year", "category", "amount_usd_mln"]
    missing = [r for r in required if r not in df.columns]
    if missing:
        raise ValueError(f"Missing columns for costs: {missing}. Available: {list(df.columns)}")

    if "is_oil" not in df.columns and "fluid_type" in df.columns:
        df["is_oil"] = df["fluid_type"].str.lower().str.contains("oil", na=False)
    if "is_gas" not in df.columns and "fluid_type" in df.columns:
        df["is_gas"] = df["fluid_type"].str.lower().str.contains("gas", na=False)

    df["is_oil"] = df.get("is_oil", True)
    df["is_gas"] = df.get("is_gas", False)

    items = []
    for _, row in df.iterrows():
        try:
            item = CostItem(
                block_name=str(row["block_name"]).strip(),
                year=int(row["year"]),
                category=str(row["category"]).lower().replace(" ", "_"),
                amount_usd_mln=float(row["amount_usd_mln"]),
                is_oil=bool(row.get("is_oil", True)),
                is_gas=bool(row.get("is_gas", False)),
            )
            items.append(item)
        except Exception as e:
            print(f"Skipping bad cost row: {e}")
    return items


def _load_costs_from_workbook(wb, fluid: str = "oil") -> list[CostProfile]:
    """Generalized EEM cost load from open workbook."""
    sheet = "Block_TC" if fluid == "oil" else "Block_TC_Gas"
    if sheet not in wb.sheetnames:
        return []
    ws = wb[sheet]

    # Use similar discovery as before
    block_starts = []
    for c in range(2, ws.max_column + 1):
        cat = ws.cell(row=2, column=c).value
        if cat == "Exploration":
            for look in [0, -1, -2, 2, 3]:
                candidate = ws.cell(row=1, column=max(1, c + look)).value
                if candidate and isinstance(candidate, str) and ("OML" in candidate or "FIELD" not in candidate.upper()):
                    block_starts.append((c, candidate.strip()))
                    break
    if not block_starts:
        for c in range(3, ws.max_column + 1, 6):
            nm = ws.cell(row=1, column=c).value
            if nm and isinstance(nm, str) and "OML" in nm:
                block_starts.append((c, nm.strip()))

    years = []
    for r in range(4, ws.max_row + 1):
        y = ws.cell(row=r, column=1).value
        if isinstance(y, (int, float)) and 2000 < y < 2100:
            years.append(int(y))

    profiles = []
    for start_col, name in block_starts:
        items = []
        for r, year in enumerate(years, start=4):
            for offset, excel_cat in enumerate(["Exploration", "CAPEX Wells", "CAPEX Facilities", "OPEX"]):
                col = start_col + offset
                amt = ws.cell(row=r, column=col).value or 0.0
                try:
                    amt = float(amt)
                except:
                    amt = 0.0
                model_cat = CATEGORY_MAP.get(excel_cat, "other")
                item = CostItem(
                    block_name=name,
                    year=year,
                    category=model_cat,
                    amount_usd_mln=amt,
                    is_oil=(fluid == "oil"),
                    is_gas=(fluid == "gas"),
                )
                items.append(item)
        profiles.append(CostProfile(block_name=name, items=items))
    return profiles


def import_costs_from_file(
    source: Union[Path, str, UploadedFileLike],
    eem_structure: bool = True,
    column_map: Optional[dict[str, str]] = None,
    fluid: str = "oil",
    **kwargs: Any,
) -> list[CostProfile]:
    """Import cost data from various formats, similar to production.

    For EEM files uses specialized parser.
    For generic tabular uses _cost_items_from_dataframe and groups into CostProfiles.
    """
    if hasattr(source, "name"):
        ext = Path(getattr(source, "name", "file")).suffix.lower()
        is_upload = True
    else:
        p = Path(str(source))
        ext = p.suffix.lower()
        is_upload = False

    print(f"[costs_handler] import_costs_from_file: {source} ext={ext}")

    if ext in {".xlsx", ".xlsm"}:
        if is_upload:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(source.getvalue() if hasattr(source, "getvalue") else source.read())
                tmp_path = tmp.name
            try:
                wb = load_workbook(tmp_path, data_only=True)
            finally:
                Path(tmp_path).unlink(missing_ok=True)
        else:
            wb = load_workbook(str(source), data_only=True)
        try:
            if eem_structure and "Block_TC" in wb.sheetnames:
                return _load_costs_from_workbook(wb, fluid=fluid)
            else:
                # generic sheet to df
                sheet = kwargs.get("sheet_name") or wb.sheetnames[0]
                data = list(wb[sheet].values)
                df = pd.DataFrame(data[1:], columns=data[0]) if data else pd.DataFrame()
                items = _cost_items_from_dataframe(df, column_map=column_map)
                # group into profiles
                by_block = {}
                for item in items:
                    by_block.setdefault(item.block_name, []).append(item)
                return [CostProfile(block_name=k, items=v) for k, v in by_block.items()]
        finally:
            wb.close()

    elif ext in {".csv", ".txt", ".tsv"}:
        sep = "\t" if ext == ".tsv" else None
        if is_upload:
            df = pd.read_csv(source, sep=sep)
        else:
            df = pd.read_csv(source, sep=sep)
        items = _cost_items_from_dataframe(df, column_map=column_map)
        by_block = {}
        for item in items:
            by_block.setdefault(item.block_name, []).append(item)
        return [CostProfile(block_name=k, items=v) for k, v in by_block.items()]

    elif ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Install python-docx for .docx support")
        if is_upload:
            content = source.getvalue() if hasattr(source, "getvalue") else source.read()
            doc = Document(io.BytesIO(content))
        else:
            doc = Document(str(source))
        if not doc.tables:
            return []
        table = doc.tables[0]
        data = [[c.text for c in row.cells] for row in table.rows]
        df = pd.DataFrame(data[1:], columns=data[0]) if len(data) > 1 else pd.DataFrame()
        items = _cost_items_from_dataframe(df, column_map=column_map)
        by_block = {}
        for item in items:
            by_block.setdefault(item.block_name, []).append(item)
        return [CostProfile(block_name=k, items=v) for k, v in by_block.items()]

    elif ext == ".xls":
        try:
            df = pd.read_excel(source if not is_upload else source, sheet_name=kwargs.get("sheet_name", 0))
            items = _cost_items_from_dataframe(df, column_map=column_map)
            by_block = {}
            for item in items:
                by_block.setdefault(item.block_name, []).append(item)
            return [CostProfile(block_name=k, items=v) for k, v in by_block.items()]
        except Exception as e:
            raise ValueError(f".xls limited support: convert to xlsx. {e}")

    else:
        raise ValueError(f"Unsupported for costs: {ext}")


def import_costs_from_uploaded_file(uploaded: UploadedFileLike, **kwargs) -> list[CostProfile]:
    return import_costs_from_file(uploaded, **kwargs)
