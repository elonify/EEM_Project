"""
src/EEM_Project/production/data_handler.py

Data input layer for production profiles.

Supports two modes (per architecture):
- Import / load from the source xlsm (wide Block_Oil Data + Block_Gas Data format).
- Load from previously exported JSON (for offline use or after manual edits).
- In-app editing will use the same ProductionProfile models + roundtrip to JSON/ later DB.

All loaded data validates against the Pydantic models.

Provides:
- load_production_profiles()
- save_profiles()
- to_summary_dataframe()  (for UI tables, validation against Prod_Summary)
- filter_by_block(), get_blocks(), etc.

The handler is the single source for "what production data are we using in this run".
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Optional

import pandas as pd
from openpyxl import load_workbook

from EEM_Project.core.models import (
    GasUtilization,
    ProductionProfile,
    TerrainType,
)

# Default external source (update if your working copy moves)
DEFAULT_XLSM = Path(r"C:\Users\Emmanuel Onwuka\Desktop\Acquisitions\Econ_Model_Draft_2.xlsm")


def _discover_blocks(ws) -> list[tuple[int, str]]:
    """Find (daily_col, block_name) pairs from row 1 of Block_* Data sheets."""
    blocks = []
    # Row 1 has "FIELD NAME:" in A, then block names every 3 columns starting at B (col 2)
    for c in range(2, ws.max_column + 1, 3):
        val = ws.cell(row=1, column=c).value
        if val and isinstance(val, str):
            name = val.strip()
            # Include core fields, contingents, and useful totals/scenarios
            if name and ("OML" in name or "FIELD" not in name.upper()):
                blocks.append((c, name))
    return blocks


def load_production_profiles(
    xlsm_path: Optional[Path] = None,
    json_dir: Optional[Path] = None,
    include_totals: bool = False,
) -> list[ProductionProfile]:
    """
    Primary entry point for data input (backward compatible).

    Now also supports many more formats via the new `import_from_file` / `import_from_uploaded_file` functions.

    Priority:
    1. If xlsm_path (or default) exists and is readable -> parse the wide sheets (preferred, always fresh).
    2. Else fall back to loading all *.json from json_dir (e.g. data/examples/ after a previous --full import).

    Returns list of ProductionProfile (oil + gas for each included block).
    """
    if xlsm_path is None:
        xlsm_path = DEFAULT_XLSM

    if xlsm_path and xlsm_path.exists():
        return _load_from_xlsm(xlsm_path, include_totals=include_totals)

    if json_dir is None:
        json_dir = Path("data/examples")

    if json_dir.exists():
        return _load_from_json_dir(json_dir)

    raise FileNotFoundError(
        f"No xlsm at {xlsm_path} and no JSON profiles in {json_dir}. "
        "Run scripts/import_from_xlsm.py --full first, or provide a valid --path / xlsm_path."
    )


def _load_from_xlsm(xlsm_path: Path, include_totals: bool = False) -> list[ProductionProfile]:
    print(f"[data_handler] Loading production directly from xlsm: {xlsm_path}")
    wb = load_workbook(str(xlsm_path), data_only=True)

    # Get consistent year list from Block_TC
    years = _get_years(wb["Block_TC"])

    profiles: list[ProductionProfile] = []

    for sheet_name, fluid in [("Block_Oil Data", "oil"), ("Block_Gas Data", "gas")]:
        ws = wb[sheet_name]
        blocks = _discover_blocks(ws)

        for daily_col, name in blocks:
            if not include_totals and ("TOTAL" in name.upper() or "SCENARIO" in name.upper() or "COPIED" in name.upper()):
                continue  # skip aggregate rows by default; user can opt-in

            annual_col = daily_col + 1
            daily_rates: list[float] = []
            annual_vols: list[float] = []

            for r in range(4, 4 + len(years)):
                d = ws.cell(row=r, column=daily_col).value
                a = ws.cell(row=r, column=annual_col).value
                daily_rates.append(float(d) if isinstance(d, (int, float)) else 0.0)
                annual_vols.append(float(a) if isinstance(a, (int, float)) else 0.0)

            # Infer some metadata from name / global Ec_IO if needed
            oml = "OML 123" if "OML 123" in name else ("OML 20" if "OML 20" in name else None)
            terrain = TerrainType.SHALLOW_WATER  # Currently driven by Ec_IO G20 for all runs
            gas_util = GasUtilization.DOMESTIC

            prof = ProductionProfile(
                block_name=name,
                oml=oml,
                fluid_type=fluid,  # type: ignore[arg-type]
                terrain=terrain,
                gas_utilization=gas_util,
                years=years,
                daily_rates_kbd=daily_rates,
                annual_volumes=annual_vols,
                source=f"{sheet_name} - {name}",
            )
            profiles.append(prof)

    wb.close()
    print(f"[data_handler] Loaded {len(profiles)} production profiles (oil+gas).")
    return profiles


def _get_years(ws) -> list[int]:
    years: list[int] = []
    for r in range(4, ws.max_row + 1):
        y = ws.cell(row=r, column=1).value
        if isinstance(y, (int, float)) and 2000 < y < 2100:
            years.append(int(y))
    return years


def _load_from_json_dir(json_dir: Path) -> list[ProductionProfile]:
    profiles = []
    for jf in sorted(json_dir.glob("*.json")):
        if "oil" in jf.name or "gas" in jf.name:  # our export convention
            with open(jf, "r", encoding="utf-8") as f:
                data = json.load(f)
            prof = ProductionProfile(**data)
            profiles.append(prof)
    print(f"[data_handler] Loaded {len(profiles)} profiles from JSON dir {json_dir}")
    return profiles


def save_profiles(profiles: list[ProductionProfile], out_dir: Path) -> None:
    """Export list of profiles to individual JSON files (one per block+fluid)."""
    out_dir.mkdir(parents=True, exist_ok=True)
    for prof in profiles:
        safe_name = prof.block_name.replace(" ", "_").replace("/", "_").replace(":", "").replace("-", "_")
        out_file = out_dir / f"{safe_name}_{prof.fluid_type}.json"
        with open(out_file, "w", encoding="utf-8") as f:
            json.dump(prof.model_dump(), f, indent=2, default=str)
    print(f"[data_handler] Saved {len(profiles)} profiles to {out_dir}")


def to_summary_dataframe(profiles: list[ProductionProfile], fluid: Optional[str] = None) -> pd.DataFrame:
    """
    Long-format DataFrame suitable for st.data_editor, charts, and comparison to Prod_Summary sheet.
    Columns: block_name, fluid_type, year, daily_rate, annual_volume
    """
    rows = []
    for p in profiles:
        if fluid and p.fluid_type != fluid:
            continue
        for i, yr in enumerate(p.years):
            rows.append({
                "block_name": p.block_name,
                "fluid_type": p.fluid_type,
                "year": yr,
                "daily_rate": p.daily_rates_kbd[i] if i < len(p.daily_rates_kbd) else 0.0,
                "annual_volume": p.annual_volumes[i] if i < len(p.annual_volumes) else 0.0,
            })
    return pd.DataFrame(rows)


def get_unique_blocks(profiles: list[ProductionProfile]) -> list[str]:
    return sorted({p.block_name for p in profiles})


# Convenience for UI / tests
def load_core_oml123_profiles(xlsm_path: Optional[Path] = None) -> list[ProductionProfile]:
    """Load only the main non-contingent, non-total blocks for OML 123 focus."""
    profs = load_production_profiles(xlsm_path=xlsm_path, include_totals=False)
    core_keywords = ["INAGHA", "ORON WEST C -", "ADANGA SOUTH", "ADANGA MAIN -", "ADANGA WEST -", "ADNH -", "NORTH ORON", "EBUGHU MAIN", "EBUGHU NORTH EAST", "KITA MARINE -"]
    return [p for p in profs if any(kw in p.block_name for kw in core_keywords)]


# ============================================================
# NEW: Flexible multi-format import support
# ============================================================

import io
import sys
import tempfile
from typing import Any, BinaryIO, TextIO, Union

UploadedFileLike = Union[BinaryIO, "UploadedFile"]  # for type hints with Streamlit


def _profiles_from_dataframe(
    df: pd.DataFrame, column_map: Optional[dict[str, str]] = None, **kwargs
) -> list[ProductionProfile]:
    """Convert a generic tabular DataFrame into ProductionProfile objects.

    Supports common column name aliases. All data will be validated by the Pydantic model.
    """
    if column_map:
        df = df.rename(columns=column_map)

    # Auto-detect common column names (case-insensitive)
    col_map = {}
    aliases = {
        "block_name": ["block_name", "block", "field", "well", "block/field", "name", "block name", "field name"],
        "year": ["year", "yr", "date", "period", "time", "years"],
        "daily_rate": ["daily_rate", "daily", "daily production", "mb/d", "daily oil", "daily gas", "rate"],
        "annual_volume": ["annual_volume", "annual", "volume", "annual production", "mmbbl", "bscf", "annual oil"],
        "fluid_type": ["fluid_type", "fluid", "type", "fluid type", "oil/gas", "product"],
    }

    for std_name, possible in aliases.items():
        for col in df.columns:
            col_lower = str(col).lower().strip()
            if col_lower in possible:
                col_map[col] = std_name
                break

    df = df.rename(columns=col_map)

    # Ensure required columns
    required = ["block_name", "year", "annual_volume"]
    missing = [r for r in required if r not in df.columns]
    if missing:
        raise ValueError(
            f"Could not find required columns {missing}. "
            f"Available: {list(df.columns)}. "
            "Please use column_map={'your_col': 'block_name', ...} or rename your columns."
        )

    if "fluid_type" not in df.columns:
        df["fluid_type"] = kwargs.get("default_fluid", "oil")

    if "daily_rate" not in df.columns:
        df["daily_rate"] = 0.0

    # Clean types
    df = df.dropna(subset=["block_name", "year"])  # at minimum
    df["year"] = pd.to_numeric(df["year"], errors="coerce").astype("Int64")
    df = df.dropna(subset=["year"])

    profiles: list[ProductionProfile] = []
    for (bname, ftype), g in df.groupby(["block_name", "fluid_type"], dropna=False):
        g = g.sort_values("year")
        prof = ProductionProfile(
            block_name=str(bname).strip(),
            fluid_type=str(ftype).strip().lower(),
            years=[int(y) for y in g["year"].tolist()],
            daily_rates_kbd=[float(x) if pd.notna(x) else 0.0 for x in g["daily_rate"].tolist()],
            annual_volumes=[float(x) if pd.notna(x) else 0.0 for x in g["annual_volume"].tolist()],
            terrain=TerrainType.SHALLOW_WATER,
            gas_utilization=GasUtilization.DOMESTIC,
            source=f"Imported from file - {bname}",
        )
        profiles.append(prof)

    print(f"[data_handler] Imported {len(profiles)} profiles from generic DataFrame.")
    return profiles


def _load_from_workbook(wb, include_totals: bool = False) -> list[ProductionProfile]:
    """Internal: load EEM wide format from an already open openpyxl workbook."""
    years = _get_years(wb["Block_TC"])

    profiles: list[ProductionProfile] = []

    for sheet_name, fluid in [("Block_Oil Data", "oil"), ("Block_Gas Data", "gas")]:
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        blocks = _discover_blocks(ws)

        for daily_col, name in blocks:
            if not include_totals and any(x in name.upper() for x in ["TOTAL", "SCENARIO", "COPIED"]):
                continue

            annual_col = daily_col + 1
            daily_rates: list[float] = []
            annual_vols: list[float] = []

            for r in range(4, 4 + len(years)):
                d = ws.cell(row=r, column=daily_col).value
                a = ws.cell(row=r, column=annual_col).value
                daily_rates.append(float(d) if isinstance(d, (int, float)) else 0.0)
                annual_vols.append(float(a) if isinstance(a, (int, float)) else 0.0)

            prof = ProductionProfile(
                block_name=name,
                oml="OML 123" if "OML 123" in name else None,
                fluid_type=fluid,
                terrain=TerrainType.SHALLOW_WATER,
                gas_utilization=GasUtilization.DOMESTIC,
                years=years,
                daily_rates_kbd=daily_rates,
                annual_volumes=annual_vols,
                source=f"{sheet_name} - {name}",
            )
            profiles.append(prof)

    return profiles


def import_from_file(
    source: Union[Path, str, UploadedFileLike],
    eem_structure: bool = True,
    column_map: Optional[dict[str, str]] = None,
    **kwargs: Any,
) -> list[ProductionProfile]:
    """Import production data from a wide variety of file formats.

    Supported:
      - .xlsx, .xlsm : EEM wide structure (if eem_structure=True) or generic table
      - .csv, .txt, .tsv : generic tabular data (pandas)
      - .docx : extracts tables (first table used by default)
      - .xls : limited support (convert recommended)

    For generic files, use column_map to map your columns, e.g.:
        column_map={"My Block Col": "block_name", "Yr": "year", "Vol": "annual_volume"}

    The function returns validated ProductionProfile objects ready for the UI,
    editing, and calculations.

    This makes data input extremely dynamic.
    """
    # Normalize source
    if hasattr(source, "name"):  # Streamlit UploadedFile or similar
        name = getattr(source, "name", "uploaded_file")
        ext = Path(name).suffix.lower()
        is_upload = True
    else:
        p = Path(str(source))
        ext = p.suffix.lower()
        name = str(p)
        is_upload = False

    print(f"[data_handler] import_from_file: {name} (ext={ext}, eem_structure={eem_structure})")

    # Excel family
    if ext in {".xlsx", ".xlsm"}:
        if is_upload:
            # Save upload to temp file (most reliable for openpyxl + data_only)
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(source.getvalue() if hasattr(source, "getvalue") else source.read())
                tmp_path = tmp.name
            try:
                wb = load_workbook(tmp_path, data_only=True)
            finally:
                Path(tmp_path).unlink(missing_ok=True)
        else:
            wb = load_workbook(str(source), data_only=True)

        try:
            if eem_structure and any("Block_Oil Data" in s for s in wb.sheetnames):
                return _load_from_workbook(wb, include_totals=kwargs.get("include_totals", False))
            else:
                # Generic Excel
                sheet = kwargs.get("sheet_name") or wb.sheetnames[0]
                ws = wb[sheet]
                data = list(ws.values)
                if not data:
                    return []
                header = data[0]
                rows = data[1:]
                df = pd.DataFrame(rows, columns=header)
                return _profiles_from_dataframe(df, column_map=column_map)
        finally:
            wb.close()

    # Delimited text
    elif ext in {".csv", ".txt", ".tsv"}:
        sep = kwargs.get("sep")
        if ext == ".tsv" and not sep:
            sep = "\t"
        if is_upload:
            df = pd.read_csv(source, sep=sep, **{k: v for k, v in kwargs.items() if k not in ["sep"]})
        else:
            df = pd.read_csv(source, sep=sep, **kwargs)
        return _profiles_from_dataframe(df, column_map=column_map)

    # Old .xls
    elif ext == ".xls":
        try:
            if is_upload:
                df = pd.read_excel(source, sheet_name=kwargs.get("sheet_name", 0))
            else:
                df = pd.read_excel(source, sheet_name=kwargs.get("sheet_name", 0))
            return _profiles_from_dataframe(df, column_map=column_map)
        except Exception as e:
            raise ValueError(
                f".xls import failed (xlrd/calamine may be needed or file is very old). "
                f"Strongly recommend converting to .xlsx first. Original error: {e}"
            )

    # Word documents
    elif ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for .docx support. "
                "Add it with: pip install python-docx"
            )

        if is_upload:
            # UploadedFile -> bytes
            content = source.getvalue() if hasattr(source, "getvalue") else source.read()
            doc = Document(io.BytesIO(content))
        else:
            doc = Document(str(source))

        if not doc.tables:
            raise ValueError("No tables found in the .docx file.")

        # Use first table by default (user can specify table_index later)
        table_idx = kwargs.get("table_index", 0)
        table = doc.tables[table_idx]
        data = [[cell.text.strip() for cell in row.cells] for row in table.rows]

        if len(data) < 2:
            return []

        df = pd.DataFrame(data[1:], columns=data[0])
        return _profiles_from_dataframe(df, column_map=column_map)

    elif ext == ".doc":
        raise NotImplementedError(
            ".doc (old Word format) is not directly supported. Please convert to .docx or export tables to Excel/CSV."
        )

    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            "Supported: .xlsx, .xlsm, .csv, .txt, .tsv, .docx (and limited .xls)."
        )


# Convenience wrapper for Streamlit file_uploader
def import_from_uploaded_file(
    uploaded_file: UploadedFileLike, **kwargs
) -> list[ProductionProfile]:
    """Wrapper specifically for st.file_uploader results."""
    return import_from_file(uploaded_file, **kwargs)
